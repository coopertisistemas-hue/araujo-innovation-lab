from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "executive" / "executive-overview.md"
OUT_DIR = ROOT / "output" / "doc" / "master-executive-google-docs-final"
ASSETS_DIR = OUT_DIR / "assets"
OUTPUT_DOCX = OUT_DIR / "araujo-innovation-lab-documento-master-executivo-final.docx"


COLORS = {
    "bg": "#F7F7F5",
    "panel": "#FFFFFF",
    "blue": "#1F3A45",
    "blue_soft": "#EAF0EE",
    "blue_soft_2": "#F1F5F3",
    "gold": "#C7A446",
    "gold_soft": "#F5EFDB",
    "ink": "#1F2C31",
    "body": "#34434A",
    "muted": "#6E7A80",
    "line": "#D7DEDA",
    "green_soft": "#EEF4F0",
}

SECTION_IMAGE_MAP = {
    "3. Visao Estrategica": ["matriz-6d-aprendizagem.png"],
    "6. Pilares Estrategicos": ["pilares-marketing-digital.png"],
    "7. Educacao, Mercado e IA em Uma Mesma Arquitetura": [
        "ecossistema-educacional.png"
    ],
    "10. Governanca, Continuidade e Viabilidade": ["governanca-educacional.png"],
    "11. Roadmap Estrategico": [
        "implementacao-progressiva.png",
        "mvp-operacional.png",
    ],
}

SECTION_BREAKS_BEFORE = {
    "5. Tendencias e Leitura de Futuro",
    "8. Transformacao Digital Institucional",
    "11. Roadmap Estrategico",
    "13. Visao Futura e Replicabilidade",
}


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def rgb(hex_value: str) -> tuple[int, int, int]:
    hex_value = hex_value.lstrip("#")
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def font_candidates(name: str) -> list[Path]:
    win = Path("C:/Windows/Fonts")
    mapping = {
        "serif": [win / "georgia.ttf", win / "times.ttf"],
        "serif_bold": [win / "georgiab.ttf", win / "timesbd.ttf"],
        "sans": [win / "arial.ttf", win / "calibri.ttf"],
        "sans_bold": [win / "arialbd.ttf", win / "calibrib.ttf"],
        "sans_italic": [win / "ariali.ttf", win / "calibrii.ttf"],
    }
    return mapping[name]


def get_font(name: str, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for candidate in font_candidates(name):
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


FONT_SERIF_44 = get_font("serif_bold", 44)
FONT_SERIF_32 = get_font("serif_bold", 32)
FONT_SERIF_24 = get_font("serif_bold", 24)
FONT_SANS_18 = get_font("sans", 18)
FONT_SANS_20 = get_font("sans", 20)
FONT_SANS_22_B = get_font("sans_bold", 22)
FONT_SANS_24_B = get_font("sans_bold", 24)
FONT_SANS_28_B = get_font("sans_bold", 28)
FONT_SANS_30_B = get_font("sans_bold", 30)
FONT_SANS_34_B = get_font("sans_bold", 34)
FONT_SANS_16 = get_font("sans", 16)
FONT_SANS_16_B = get_font("sans_bold", 16)
FONT_SANS_15 = get_font("sans", 15)
FONT_SANS_14 = get_font("sans", 14)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        box = draw.textbbox((0, 0), candidate, font=font)
        if box[2] - box[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_lines(
    draw: ImageDraw.ImageDraw,
    text: str,
    font,
    x: int,
    y: int,
    max_width: int,
    fill: str,
    line_gap: int = 6,
) -> int:
    lines = wrap_text(draw, text, font, max_width)
    current_y = y
    for line in lines:
        draw.text((x, current_y), line, font=font, fill=fill)
        box = draw.textbbox((x, current_y), line, font=font)
        current_y += (box[3] - box[1]) + line_gap
    return current_y


def create_canvas(width: int = 1600, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), rgb(COLORS["bg"]))
    draw = ImageDraw.Draw(img)
    return img, draw


def add_visual_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    draw.rounded_rectangle((70, 50, 220, 90), radius=10, fill=rgb(COLORS["blue_soft"]))
    draw.text((88, 61), "ATIVO VISUAL", font=FONT_SANS_16_B, fill=rgb(COLORS["blue"]))
    draw.text((70, 120), title, font=FONT_SERIF_44, fill=rgb(COLORS["ink"]))
    draw.text((70, 182), subtitle, font=FONT_SANS_20, fill=rgb(COLORS["muted"]))
    draw.line((70, 225, 1530, 225), fill=rgb(COLORS["line"]), width=2)


def save_image(img: Image.Image, name: str) -> Path:
    path = ASSETS_DIR / name
    img.save(path)
    return path


def generate_roadmap() -> Path:
    img, draw = create_canvas(1600, 780)
    add_visual_header(
        draw,
        "Implementacao Progressiva",
        "Leitura institucional de maturacao, controle de complexidade e continuidade.",
    )
    x_positions = [280, 800, 1320]
    titles = ["Ano 1", "Ano 2", "Ano 3"]
    subtitles = ["Estruturacao", "Maturacao", "Expansao"]
    details = [
        "Adaptacao, organizacao e cultura de laboratorio.",
        "KPIs, projetos mais estruturados e integracao gradual com mercado.",
        "Experiencias ampliadas, portfolio consolidado e maturidade operacional.",
    ]
    draw.line((180, 430, 1420, 430), fill=rgb(COLORS["blue"]), width=6)
    for i, x in enumerate(x_positions):
        fill = COLORS["gold"] if i == 1 else COLORS["blue"]
        draw.ellipse((x - 36, 394, x + 36, 466), fill=rgb(fill), outline=rgb(fill))
        draw.text((x - 36, 500), titles[i], font=FONT_SANS_24_B, fill=rgb(COLORS["ink"]))
        draw.text((x - 72, 540), subtitles[i], font=FONT_SERIF_24, fill=rgb(COLORS["blue"]))
        draw_lines(draw, details[i], FONT_SANS_18, x - 145, 590, 290, COLORS["body"], 4)
    draw.rounded_rectangle((100, 660, 1500, 730), radius=16, fill=rgb(COLORS["blue_soft_2"]))
    draw.text(
        (130, 683),
        "Principio orientador: implementacao progressiva, com prioridade para sustentabilidade institucional e qualidade pedagogica.",
        font=FONT_SANS_18,
        fill=rgb(COLORS["body"]),
    )
    return save_image(img, "implementacao-progressiva.png")


def generate_matrix() -> Path:
    img, draw = create_canvas(1600, 980)
    add_visual_header(
        draw,
        "Matriz 6D de Aprendizagem",
        "Dimensoes complementares que articulam conhecimento, experiencia e demonstracao de competencias.",
    )
    cards = [
        ("Teoria", "Base conceitual", "Fundamenta linguagem, repertorio e criterio."),
        ("Laboratorio", "Pratica guiada", "Transforma conceito em experimentacao disciplinada."),
        ("Projetos", "Entregas aplicadas", "Estrutura producao, responsabilidade e refinamento."),
        ("Experiencias", "Observacao territorial", "Conecta formacao a contextos reais e sinais de mercado."),
        ("Mercado", "Leitura aplicada", "Aproxima tendencias, canais e dinamicas contemporaneas."),
        ("Portfolio", "Evidencia de trajetoria", "Consolida demonstracoes consistentes de aprendizagem."),
    ]
    x0, y0 = 80, 280
    card_w, card_h = 450, 240
    gap_x, gap_y = 40, 35
    for idx, (title, subtitle, desc) in enumerate(cards):
        row = idx // 3
        col = idx % 3
        x = x0 + col * (card_w + gap_x)
        y = y0 + row * (card_h + gap_y)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=24, fill=rgb(COLORS["panel"]), outline=rgb(COLORS["line"]), width=2)
        draw.rounded_rectangle((x + 24, y + 24, x + 92, y + 92), radius=18, fill=rgb(COLORS["blue_soft"]))
        draw.text((x + 44, y + 42), str(idx + 1), font=FONT_SANS_22_B, fill=rgb(COLORS["blue"]))
        draw.text((x + 122, y + 28), title, font=FONT_SANS_28_B, fill=rgb(COLORS["ink"]))
        draw.text((x + 122, y + 72), subtitle, font=FONT_SERIF_24, fill=rgb(COLORS["blue"]))
        draw_lines(draw, desc, FONT_SANS_18, x + 36, y + 128, card_w - 72, COLORS["body"], 5)
    draw.rounded_rectangle((80, 830, 1520, 900), radius=18, fill=rgb(COLORS["gold_soft"]))
    draw.text(
        (110, 855),
        "Leitura institucional: a formacao ganha densidade quando essas dimensoes operam de forma articulada, progressiva e observavel.",
        font=FONT_SANS_18,
        fill=rgb(COLORS["body"]),
    )
    return save_image(img, "matriz-6d-aprendizagem.png")


def generate_governance() -> Path:
    img, draw = create_canvas(1600, 980)
    add_visual_header(
        draw,
        "Governanca Educacional",
        "Fluxo de sustentacao institucional e operacao pedagogica com clareza e baixa burocracia.",
    )
    boxes = {
        "Direcao": (620, 270, 980, 350),
        "Orientacao Pedagogica": (200, 430, 620, 520),
        "Professora Titular": (640, 430, 1040, 520),
        "Professor Orientador": (1060, 430, 1460, 520),
        "Professores Interdisciplinares": (340, 620, 760, 710),
        "Estudantes": (840, 620, 1260, 710),
        "Parceiros Externos": (1120, 800, 1490, 880),
    }
    for title, box in boxes.items():
        fill = COLORS["blue_soft"] if title in {"Direcao", "Estudantes"} else COLORS["panel"]
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=20, fill=rgb(fill), outline=rgb(COLORS["line"]), width=2)
        draw.text((x1 + 28, y1 + 26), title, font=FONT_SANS_24_B, fill=rgb(COLORS["ink"]))
    lines = [
        ((800, 350), (410, 430)),
        ((800, 350), (840, 430)),
        ((800, 350), (1260, 430)),
        ((410, 520), (550, 620)),
        ((840, 520), (1050, 620)),
        ((1260, 520), (1050, 620)),
        ((1260, 520), (1310, 800)),
    ]
    for start, end in lines:
        draw.line((*start, *end), fill=rgb(COLORS["blue"]), width=3)
    draw.rounded_rectangle((90, 800, 1030, 890), radius=18, fill=rgb(COLORS["blue_soft_2"]))
    draw.text(
        (120, 825),
        "Leitura executiva: coordenacao clara, apoio interdisciplinar, acompanhamento pedagogico e abertura controlada para cooperacoes externas.",
        font=FONT_SANS_18,
        fill=rgb(COLORS["body"]),
    )
    return save_image(img, "governanca-educacional.png")


def generate_ecosystem() -> Path:
    img, draw = create_canvas(1600, 980)
    add_visual_header(
        draw,
        "Ecossistema Educacional",
        "Transformacao educacional regional integrada entre escola, territorio, mercado e competencias contemporaneas.",
    )
    center = (800, 540)
    nodes = [
        ("Estudantes", 800, 330),
        ("Experiencias Praticas", 1035, 420),
        ("Empresas Locais", 1090, 655),
        ("Turismo", 800, 770),
        ("Comunidade", 510, 655),
        ("Marketing Digital", 470, 420),
        ("IA", 800, 210),
    ]
    draw.ellipse((590, 430, 1010, 650), fill=rgb(COLORS["blue_soft"]), outline=rgb(COLORS["blue"]), width=3)
    draw.text((675, 500), "EEB Araujo", font=FONT_SANS_34_B, fill=rgb(COLORS["ink"]))
    draw.text((650, 548), "Innovation Lab", font=FONT_SERIF_32, fill=rgb(COLORS["blue"]))
    for label, x, y in nodes:
        draw.line((center[0], center[1], x, y), fill=rgb(COLORS["line"]), width=3)
        box = (x - 140, y - 42, x + 140, y + 42)
        fill = COLORS["panel"] if label != "IA" else COLORS["gold_soft"]
        draw.rounded_rectangle(box, radius=20, fill=rgb(fill), outline=rgb(COLORS["line"]), width=2)
        tw = draw.textbbox((0, 0), label, font=FONT_SANS_22_B)
        draw.text((x - (tw[2] - tw[0]) / 2, y - 14), label, font=FONT_SANS_22_B, fill=rgb(COLORS["ink"]))
    draw.rounded_rectangle((120, 840, 1480, 910), radius=18, fill=rgb(COLORS["green_soft"]))
    draw.text(
        (150, 865),
        "A escola opera como centro articulador de aprendizagem aplicada, leitura territorial e conexao progressiva com a economia digital.",
        font=FONT_SANS_18,
        fill=rgb(COLORS["body"]),
    )
    return save_image(img, "ecossistema-educacional.png")


def generate_mvp() -> Path:
    img, draw = create_canvas(1600, 900)
    add_visual_header(
        draw,
        "MVP Operacional",
        "Camada de implementacao controlada para reduzir sobrecarga e preservar sustentabilidade institucional.",
    )
    columns = [
        (
            "Nucleo Minimo Viavel",
            [
                "Branding institucional basico",
                "Redes sociais institucionais",
                "Google Meu Negocio",
                "Copywriting introdutorio",
                "Laboratorio simplificado",
            ],
        ),
        (
            "Prioridades Iniciais",
            [
                "Portfolio leve",
                "Introducao a IA",
                "Reputacao digital",
                "1 projeto integrador",
                "Experiencias observacionais",
            ],
        ),
        (
            "Expansao Futura",
            [
                "Automacao moderada",
                "Analytics mais robusto",
                "Projetos ampliados",
                "Mais relacoes com mercado",
                "Maturidade operacional",
            ],
        ),
    ]
    start_x, y = 85, 285
    w, h, gap = 450, 420, 40
    for idx, (title, items) in enumerate(columns):
        x = start_x + idx * (w + gap)
        fill = COLORS["panel"] if idx != 1 else COLORS["blue_soft"]
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=rgb(fill), outline=rgb(COLORS["line"]), width=2)
        draw.text((x + 30, y + 28), title, font=FONT_SANS_28_B, fill=rgb(COLORS["ink"]))
        current_y = y + 90
        for item in items:
            draw.ellipse((x + 30, current_y + 8, x + 42, current_y + 20), fill=rgb(COLORS["gold"]))
            current_y = draw_lines(draw, item, FONT_SANS_18, x + 58, current_y, w - 90, COLORS["body"], 4) + 10
    draw.rounded_rectangle((85, 755, 1515, 825), radius=18, fill=rgb(COLORS["gold_soft"]))
    draw.text(
        (115, 780),
        "Principio de sustentabilidade: a continuidade institucional tem prioridade sobre expansao acelerada.",
        font=FONT_SANS_18,
        fill=rgb(COLORS["body"]),
    )
    return save_image(img, "mvp-operacional.png")


def generate_pillars() -> Path:
    img, draw = create_canvas(1600, 980)
    add_visual_header(
        draw,
        "Pilares do Marketing Digital",
        "Mapa de competencias contemporaneas organizado por estrategia, operacao e evolucao.",
    )
    groups = [
        (
            "Estrategia e Marca",
            COLORS["blue_soft"],
            ["Branding", "Marketing Digital", "Copywriting", "Presenca Digital"],
            90,
        ),
        (
            "Operacao e Relacionamento",
            COLORS["green_soft"],
            ["Reputacao Digital", "Relacionamento Digital", "Experiencia do Cliente", "E-commerce"],
            560,
        ),
        (
            "Leitura e Evolucao",
            COLORS["gold_soft"],
            ["IA", "Analytics"],
            1030,
        ),
    ]
    for title, fill, items, x in groups:
        width = 420 if x < 1030 else 480
        draw.rounded_rectangle((x, 285, x + width, 560), radius=24, fill=rgb(fill), outline=rgb(COLORS["line"]), width=2)
        draw.text((x + 26, 312), title, font=FONT_SANS_28_B, fill=rgb(COLORS["ink"]))
        current_y = 380
        for item in items:
            draw.rounded_rectangle((x + 24, current_y, x + width - 24, current_y + 48), radius=14, fill=rgb(COLORS["panel"]))
            draw.text((x + 42, current_y + 12), item, font=FONT_SANS_18, fill=rgb(COLORS["body"]))
            current_y += 60
    draw.rounded_rectangle((120, 655, 1480, 820), radius=22, fill=rgb(COLORS["panel"]), outline=rgb(COLORS["line"]), width=2)
    draw.text((150, 690), "Leitura institucional", font=FONT_SANS_24_B, fill=rgb(COLORS["blue"]))
    draw_lines(
        draw,
        "Os pilares nao operam como modulos isolados. Eles estruturam uma formacao que combina linguagem, reputacao, relacionamento, dados e uso criterioso de tecnologias contemporaneas.",
        FONT_SANS_20,
        150,
        735,
        1280,
        COLORS["body"],
        5,
    )
    return save_image(img, "pilares-marketing-digital.png")


def build_visual_assets() -> dict[str, Path]:
    return {
        "roadmap": generate_roadmap(),
        "matrix": generate_matrix(),
        "governance": generate_governance(),
        "ecosystem": generate_ecosystem(),
        "mvp": generate_mvp(),
        "pillars": generate_pillars(),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "D7DEDA", size: str = "10") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    fld.append(run)
    paragraph._p.append(fld)


def set_paragraph_bottom_border(paragraph, color: str = "D7DEDA", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run_style(run, font_name: str, size: int, color: str, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("34434A")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    title = doc.styles["Title"]
    title.font.name = "Georgia"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string("1F2C31")
    title.font.bold = True

    for style_name, size in [("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Georgia" if style_name == "Heading 1" else "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F3A45")
        style.paragraph_format.space_before = Pt(18 if style_name == "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(8 if style_name == "Heading 1" else 4)

    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
        caption.font.name = "Arial"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        caption.font.size = Pt(9)
        caption.font.color.rgb = RGBColor.from_string("6E7A80")


def configure_section(section, with_header: bool = True) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.2)

    if with_header:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.text = "Araujo Innovation Lab | Documento Master Executivo"
        for run in p.runs:
            set_run_style(run, "Arial", 9, "6E7A80")
        set_paragraph_bottom_border(p)

        footer = section.footer
        table = footer.add_table(rows=1, cols=3, width=Inches(6.0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = table.rows[0].cells
        contents = [
            "Documento Institucional | v1.0",
            "Pagina ",
            "EEB Araujo Figueiredo | Maio 2026",
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for cell, text, align in zip(cells, contents, aligns):
            p = cell.paragraphs[0]
            p.alignment = align
            run = p.add_run(text)
            set_run_style(run, "Arial", 8, "6E7A80")
            if text == "Pagina ":
                add_page_field(p)
            set_cell_borders(cell, color="FFFFFF", size="0")


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    configure_section(section, with_header=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("ARAUJO INNOVATION LAB")
    set_run_style(run, "Arial", 10, "1F3A45", bold=True)

    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(6)
    bar.paragraph_format.space_after = Pt(18)
    set_paragraph_bottom_border(bar, color="C7A446", size="16")

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(60)
    title.add_run("Documento Master Executivo")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Programa institucional de transformacao educacional aplicada")
    set_run_style(r, "Arial", 13, "6E7A80")

    institution = doc.add_paragraph()
    institution.paragraph_format.space_before = Pt(120)
    institution.paragraph_format.space_after = Pt(8)
    r = institution.add_run("EEB Araujo Figueiredo")
    set_run_style(r, "Georgia", 15, "1F2C31", bold=True)

    city = doc.add_paragraph()
    city.paragraph_format.space_after = Pt(4)
    r = city.add_run("Urubici - Santa Catarina")
    set_run_style(r, "Arial", 11, "34434A")

    nature = doc.add_paragraph()
    r = nature.add_run("Posicionamento executivo institucional")
    set_run_style(r, "Arial", 11, "34434A")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(120)
    line = spacer.add_run("Versao final para Google Docs e PDF institucional")
    set_run_style(line, "Arial", 9, "6E7A80")


def add_manual_index(doc: Document, sections: list[dict]) -> None:
    doc.add_paragraph("Sumario Executivo", style="Heading 1")
    intro = doc.add_paragraph()
    intro.add_run(
        "Estrutura organizada para leitura institucional, articulacao externa e uso em apresentacoes formais."
    )
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for section in sections:
        row = table.add_row()
        row.cells[0].width = Inches(5.6)
        row.cells[1].width = Inches(0.6)
        row.cells[0].paragraphs[0].add_run(section["title"])
        row.cells[1].paragraphs[0].add_run("—")
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in row.cells:
            set_cell_borders(cell, color="EAF0EE", size="6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_metadata_block(doc: Document, metadata: dict[str, str]) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("Instituicao", metadata.get("Instituicao", "EEB Araujo Figueiredo")),
        ("Objetivo", metadata.get("Objetivo", "")),
        ("Resultado esperado", metadata.get("Resultado esperado", "")),
    ]
    for row, (label, value) in zip(table.rows, items):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EAF0EE")
        for cell in row.cells:
            set_cell_borders(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_style(
                        run,
                        "Arial",
                        10,
                        "1F2C31" if cell == row.cells[0] else "34434A",
                        bold=cell == row.cells[0],
                    )

    doc.add_paragraph()
    block = doc.add_table(rows=1, cols=1)
    block.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = block.rows[0].cells[0]
    set_cell_shading(cell, "F5EFDB")
    set_cell_borders(cell, color="C7A446", size="12")
    p = cell.paragraphs[0]
    r1 = p.add_run("Leitura executiva: ")
    set_run_style(r1, "Arial", 11, "1F3A45", bold=True)
    r2 = p.add_run(
        "documento institucional orientado a credibilidade, clareza narrativa, integracao com mercado e continuidade de implementacao."
    )
    set_run_style(r2, "Arial", 11, "34434A")


def add_executive_block(doc: Document, label: str, text: str, fill: str = "F1F5F3") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="D7DEDA", size="10")
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label}: ")
    set_run_style(r1, "Arial", 10, "1F3A45", bold=True)
    r2 = p.add_run(text)
    set_run_style(r2, "Arial", 10, "34434A")


def add_image_with_caption(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)


def parse_markdown() -> tuple[dict[str, str], list[dict]]:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    metadata: dict[str, str] = {}
    sections: list[dict] = []
    current_section: dict | None = None
    paragraph_buffer: list[str] = []
    skip_visual = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer, current_section
        if current_section and paragraph_buffer:
            joined = " ".join(part.strip() for part in paragraph_buffer).strip()
            if joined:
                current_section["items"].append(("p", joined))
        paragraph_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("**") and ":**" in stripped and not sections:
            match = re.match(r"\*\*(.+?):\*\*\s*(.+)", stripped)
            if match:
                metadata[match.group(1)] = match.group(2)
            continue

        if stripped == "## Recomposicao Visual Recomendada":
            flush_paragraph()
            break

        if stripped == "---":
            flush_paragraph()
            skip_visual = False
            continue

        if stripped.startswith("**Sugestoes visuais para a secao**"):
            flush_paragraph()
            skip_visual = True
            continue

        if skip_visual:
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            title = stripped[3:].strip()
            if title in {
                "Araujo Innovation Lab",
            }:
                continue
            current_section = {"title": title, "items": []}
            sections.append(current_section)
            continue

        if not current_section:
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            current_section["items"].append(("h2", stripped[4:].strip()))
            continue

        if re.fullmatch(r"\*\*.+\*\*", stripped):
            flush_paragraph()
            current_section["items"].append(("h3", stripped.strip("*")))
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            current_section["items"].append(("bullet", stripped[2:].strip()))
            continue

        if not stripped:
            flush_paragraph()
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()
    return metadata, sections


def build_doc(assets: dict[str, Path], metadata: dict[str, str], sections: list[dict]) -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(main_section, with_header=True)

    add_manual_index(doc, sections)
    doc.add_paragraph()
    add_metadata_block(doc, metadata)
    doc.add_page_break()

    for section in sections:
        title = section["title"]
        if title in SECTION_BREAKS_BEFORE:
            doc.add_page_break()

        doc.add_paragraph(title, style="Heading 1")

        if title == "1. Abertura Executiva":
            add_executive_block(
                doc,
                "Mensagem central",
                "O Araujo Innovation Lab posiciona a escola como ambiente de aprendizagem aplicada, articulacao territorial e modernizacao educacional com responsabilidade institucional.",
                fill="F1F5F3",
            )

        for item_type, content in section["items"]:
            if item_type == "p":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(content)
            elif item_type == "bullet":
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.35)
                r = p.add_run(f"• {content}")
                set_run_style(r, "Arial", 11, "34434A")
            elif item_type == "h2":
                doc.add_paragraph(content, style="Heading 2")
            elif item_type == "h3":
                doc.add_paragraph(content, style="Heading 3")

        if title == "6. Pilares Estrategicos":
            add_executive_block(
                doc,
                "Leitura executiva",
                "Os pilares organizam a proposta como sistema de competencias aplicadas, sem deslocar o documento para um tom tecnico ou excessivamente operacional.",
                fill="F5EFDB",
            )

        for image_name in SECTION_IMAGE_MAP.get(title, []):
            captions = {
                "matriz-6d-aprendizagem.png": "Figura institucional 1. Matriz 6D de Aprendizagem.",
                "pilares-marketing-digital.png": "Figura institucional 2. Pilares do Marketing Digital aplicados ao laboratorio.",
                "ecossistema-educacional.png": "Figura institucional 3. Ecossistema educacional e territorial do projeto.",
                "governanca-educacional.png": "Figura institucional 4. Governanca educacional com fluxo simplificado de coordenacao.",
                "implementacao-progressiva.png": "Figura institucional 5. Implementacao progressiva em tres anos.",
                "mvp-operacional.png": "Figura institucional 6. MVP operacional para implementacao controlada.",
            }
            add_image_with_caption(doc, ASSETS_DIR / image_name, captions[image_name])

        if title == "11. Roadmap Estrategico":
            add_executive_block(
                doc,
                "Principio de sustentacao",
                "A sustentabilidade institucional tem prioridade sobre expansao acelerada, com evolucao faseada e capacidade real de continuidade.",
                fill="EEF4F0",
            )

        if title == "14. Encerramento Executivo":
            add_executive_block(
                doc,
                "Conclusao institucional",
                "O documento consolida uma proposta de transformacao educacional com ambicao qualificada, escalonamento responsavel e forte aderencia ao territorio.",
                fill="F1F5F3",
            )

    doc.save(OUTPUT_DOCX)


def main() -> None:
    ensure_dirs()
    assets = build_visual_assets()
    metadata, sections = parse_markdown()
    build_doc(assets, metadata, sections)
    print(f"Documento gerado em: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
